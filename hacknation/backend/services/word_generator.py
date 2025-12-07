"""
Word document generator service for creating legal opinion documents.

# TODO: Replace hardcoded context in generate_opinion_document() with AI-generated content
# The context dictionary should be populated by analyzing the PDF documents in the folder
# using an LLM to extract relevant information and generate appropriate text.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from datetime import datetime
import uuid


def create_opinion_document(context: dict, output_path: Path) -> Path:
    """
    Create a Word document with the legal opinion template filled with context data.
    
    Args:
        context: Dictionary containing all the template variables
        output_path: Path where the document should be saved
    
    Returns:
        Path to the created document
    """
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Title - Case number
    case_number = doc.add_paragraph()
    case_number.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = case_number.add_run(context.get("case_number", ""))
    run.bold = True
    run.font.size = Pt(14)
    
    # Main title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("OPINIA W SPRAWIE PRAWNEJ KWALIFIKACJI WYPADKU")
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    # Injured person
    doc.add_paragraph("Nazwisko i imię poszkodowanego:")
    injured = doc.add_paragraph()
    run = injured.add_run(context.get("injured_person", ""))
    run.bold = True
    
    doc.add_paragraph()
    
    # Issue description
    doc.add_paragraph("Kwestia do rozstrzygnięcia:")
    doc.add_paragraph(context.get("issue_description", ""))
    
    doc.add_paragraph()
    
    # Accident type section
    accident_date = context.get("accident_date", "")
    doc.add_paragraph(f"Czy uznać zdarzenie z dnia {accident_date} r. za wypadek podczas:")
    
    # Bullet points for accident types
    accident_types = [
        "wykonywania zwykłych czynności związanych z prowadzeniem pozarolniczej działalności",
        "wykonywania zwykłych czynności związanych ze współpracą przy prowadzeniu pozarolniczej działalności",
        "wykonywania pracy na podstawie umowy uaktywniającej",
        "opieki nad dziećmi w wieku do lat 3"
    ]
    
    for item in accident_types:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)
    
    doc.add_paragraph()
    doc.add_paragraph("lub w drodze do lub z miejsca:")
    
    travel_types = [
        "wykonywania pozarolniczej działalności",
        "współpracy przy prowadzeniu pozarolniczej działalności",
        "wykonywania pracy na podstawie umowy uaktywniającej"
    ]
    
    for item in travel_types:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)
    
    doc.add_paragraph()
    
    # Conclusion
    conclusion_header = doc.add_paragraph()
    run = conclusion_header.add_run("Wniosek:")
    run.bold = True
    doc.add_paragraph(context.get("conclusion", ""))
    
    doc.add_paragraph()
    
    # Justification
    justification_header = doc.add_paragraph()
    run = justification_header.add_run("Uzasadnienie:")
    run.bold = True
    doc.add_paragraph(context.get("justification", ""))
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Specialist signature
    specialist = doc.add_paragraph()
    specialist.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = specialist.add_run("STARSZY SPECJALISTA")
    run.bold = True
    
    specialist_name = doc.add_paragraph()
    specialist_name.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    specialist_name.add_run(context.get("specialist_name", ""))
    
    specialist_date = doc.add_paragraph()
    specialist_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    specialist_date.add_run(context.get("specialist_signature_date", ""))
    
    # Separator
    doc.add_paragraph("─" * 70)
    
    # Approbant opinion section
    _add_opinion_section(
        doc, 
        "OPINIA OSOBY UPRAWNIONEJ DO APROBATY",
        context.get("approbant_opinion", ""),
        context.get("approbant_date", ""),
        context.get("approbant_signature", "")
    )
    
    # Separator
    doc.add_paragraph("─" * 70)
    
    # Superapprobation opinion section
    _add_opinion_section(
        doc,
        "OPINIA OSOBY UPRAWNIONEJ DO SUPERAPROBATY",
        context.get("superapprobation_opinion", ""),
        context.get("superapprobation_date", ""),
        context.get("superapprobation_signature", "")
    )
    
    # Separator
    doc.add_paragraph("─" * 70)
    
    # Consultant opinion section
    _add_opinion_section(
        doc,
        "OPINIA KONSULTANTA",
        context.get("consultant_opinion", ""),
        context.get("consultant_date", ""),
        context.get("consultant_signature", "")
    )
    
    # Separator
    doc.add_paragraph("─" * 70)
    
    # Deputy director opinion section
    _add_opinion_section(
        doc,
        "OPINIA Z-CE DYREKTORA DS. ŚWIADCZEŃ",
        context.get("deputy_director_opinion", ""),
        context.get("deputy_director_date", ""),
        context.get("deputy_director_signature", "")
    )
    
    # Separator
    doc.add_paragraph("─" * 70)
    
    # Final decision section
    _add_opinion_section(
        doc,
        "DECYZJA OSOBY UPRAWNIONEJ DO SUPERAPROBATY",
        context.get("final_decision", ""),
        context.get("final_decision_date", ""),
        context.get("final_decision_signature", "")
    )
    
    # Save the document
    doc.save(str(output_path))
    return output_path


def _add_opinion_section(doc, title: str, opinion: str, date: str, signature: str):
    """Helper function to add an opinion section with title, content, date, and signature."""
    header = doc.add_paragraph()
    run = header.add_run(title)
    run.bold = True
    
    doc.add_paragraph(opinion if opinion else "")
    
    date_para = doc.add_paragraph()
    date_para.add_run(f"Data: {date}")
    
    sig_para = doc.add_paragraph()
    sig_para.add_run(f"Podpis: {signature}")
    
    doc.add_paragraph()


def generate_opinion_document(folder_path: Path, folder_name: str) -> dict:
    """
    Generate an opinion document for a given PESEL folder.
    
    # TODO: Replace hardcoded context with AI-generated content
    # This function should:
    # 1. Read all PDF files in the folder
    # 2. Extract text using OCR/pdf parsing
    # 3. Send to LLM to analyze and extract relevant information
    # 4. Generate appropriate opinion text based on the analysis
    
    Args:
        folder_path: Path to the PESEL folder containing PDF documents
        folder_name: Name of the folder (usually PESEL number)
    
    Returns:
        Dictionary with success status and generated filename
    """
    # Hardcoded example context - TODO: Replace with AI-generated content
    context = {
        "case_number": f"Znak sprawy: {folder_name}/2025",
        "injured_person": "Jan Kowalski",
        "issue_description": "Ocena czy zdarzenie spełnia definicję wypadku podczas wykonywania czynności związanych z prowadzeniem działalności gospodarczej.",
        "accident_date": "28-02-2025",
        "conclusion": "Proponuję uznać zdarzenie za wypadek podczas wykonywania czynności związanych z działalnością.",
        "justification": "Na podstawie analizy dokumentacji zgromadzonej w sprawie, w szczególności protokołu powypadkowego, zeznań świadków oraz dokumentacji medycznej, stwierdzono, że zdarzenie z dnia 28-02-2025 r. spełnia definicję wypadku w rozumieniu przepisów ustawy o ubezpieczeniu społecznym z tytułu wypadków przy pracy i chorób zawodowych.",
        "specialist_name": "Starszy Specjalista ZUS",
        "specialist_signature_date": datetime.now().strftime("%d %B %Y").upper(),
        "approbant_opinion": "Zgadzam się z opinią specjalisty. Dokumentacja jest kompletna.",
        "approbant_date": datetime.now().strftime("%d %B %Y").upper(),
        "approbant_signature": "Kierownik Wydziału",
        "superapprobation_opinion": "Akceptuję proponowaną kwalifikację prawną zdarzenia.",
        "superapprobation_date": datetime.now().strftime("%d %B %Y").upper(),
        "superapprobation_signature": "Naczelnik Wydziału",
        "consultant_opinion": "",
        "consultant_date": "",
        "consultant_signature": "",
        "deputy_director_opinion": "",
        "deputy_director_date": "",
        "deputy_director_signature": "",
        "final_decision": "",
        "final_decision_date": "",
        "final_decision_signature": ""
    }
    
    # Generate unique filename
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"Opinia_prawna_{timestamp}.docx"
    output_path = folder_path / filename
    
    try:
        create_opinion_document(context, output_path)
        return {
            "success": True,
            "filename": filename,
            "path": str(output_path)
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }

